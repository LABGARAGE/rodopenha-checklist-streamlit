#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de Check-list (DOCX/PDF) – Rodopenha

Recursos:
- Cabeçalho com LOGO (opcional via --logo)
- Nome automático: CHECKLIST_{data}_{rota}_{id}.docx
- Pasta de saída (--saida_dir) e nome customizado (--saida)
- QR Code no rodapé (--qrcode "texto/url")
- Exportação opcional para PDF (--pdf) via docx2pdf (requer MS Word no Windows/Mac)

Exemplos:
  python gerar_docx.py --id 123 --doca 05 --conferente "Maria" --turno "1º" --motorista "João" --placa "ABC1D23" \
    --destino "Rota 12" --logo "LogoWebRodoPenha-removebg-preview.png" \
    --saida_dir "saidas" --qrcode "SharePoint://RP/Checklists" --pdf \
    --notas "3524...,NFe 445654,NFe 1234/2025"

  # Modo leitor de código de barras (uma por linha; finalize com Ctrl+Z e Enter no Windows):
  python gerar_docx.py --id 123 --doca 05 --conferente "Maria" --turno "1º" --motorista "João" --placa "ABC1D23" --destino "Rota 12"
"""
import argparse
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

try:
    import qrcode
except Exception:
    qrcode = None


def add_heading_with_rule(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p_rule._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_qrcode_footer(doc, text):
    if not text or qrcode is None:
        return
    import io
    img = qrcode.make(text)
    bio = io.BytesIO(); img.save(bio, format='PNG'); bio.seek(0)
    p = doc.add_paragraph(); run = p.add_run()
    try:
        run.add_picture(bio, width=Inches(0.9))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        pass


def make_docx(output_path, meta, notas, logo_file=None, qrcode_text=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Cabeçalho: logo (esq) + título (dir)
    try:
        if logo_file and Path(logo_file).exists():
            table_hdr = doc.add_table(rows=1, cols=2)
            cell_logo = table_hdr.cell(0,0)
            cell_title = table_hdr.cell(0,1)
            run_img = cell_logo.paragraphs[0].add_run()
            run_img.add_picture(str(logo_file), width=Inches(1.2))
            pr = cell_title.paragraphs[0]
            run = pr.add_run("CHECK-LIST DE MOVIMENTACAO – RODOPENHA TRANSPORTES")
            run.bold = True; run.font.size = Pt(14)
            pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            add_heading_with_rule(doc, "CHECK-LIST DE MOVIMENTACAO – RODOPENHA TRANSPORTES")
    except Exception:
        add_heading_with_rule(doc, "CHECK-LIST DE MOVIMENTACAO – RODOPENHA TRANSPORTES")

    labels = [
        ("ID:", meta.get("id",""), "Data:", meta.get("data", datetime.now().strftime("%d/%m/%Y"))),
        ("Doca:", meta.get("doca",""), "Turno:", meta.get("turno","")),
        ("Conferente:", meta.get("conferente",""), "Motorista:", meta.get("motorista","")),
        ("Placa:", meta.get("placa",""), "Destino/Rota:", meta.get("destino","")),
    ]
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    for r, row in enumerate(labels):
        for c in range(0, 4, 2):
            cell_label = table.cell(r, c)
            cell_value = table.cell(r, c+1)
            run_l = cell_label.paragraphs[0].add_run(row[c]); run_l.bold = True
            cell_value.paragraphs[0].add_run(row[c+1])

    doc.add_paragraph("")
    p = doc.add_paragraph(); run = p.add_run("Notas Fiscais / Chaves (bipar ou digitar):"); run.bold = True

    notas = [n.strip() for n in notas if n.strip()]
    total = len(notas)
    if total == 0:
        doc.add_paragraph("—")
    else:
        cols = 2
        rows = (total + cols - 1)//cols
        t2 = doc.add_table(rows=rows, cols=cols); t2.style = "Light Grid"
        for idx, nf in enumerate(notas):
            r = idx % rows; c = idx // rows
            t2.cell(r, c).text = nf

    doc.add_paragraph("")
    t3 = doc.add_table(rows=3, cols=2); t3.style = "Table Grid"
    t3.cell(0,0).text = "Total de Notas Recebidas:"; t3.cell(0,1).text = str(total)
    t3.cell(1,0).text = "Total de Notas Expedidas:"; t3.cell(1,1).text = ""
    t3.cell(2,0).text = "Divergencias:"; t3.cell(2,1).text = ""

    doc.add_paragraph("")
    t4 = doc.add_table(rows=2, cols=2); t4.style = "Table Grid"
    t4.cell(0,0).text = "\n\nAssinatura Conferente:"; t4.cell(0,1).text = "\n\nAssinatura Motorista:"
    t4.cell(1,0).text = "Nome legivel:"; t4.cell(1,1).text = "Nome legivel:"
    doc.add_paragraph("")
    foot = doc.add_paragraph("Documento gerado automaticamente para controle interno de movimentacao.")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER; foot.runs[0].font.size = Pt(8)

    # QR Code no rodapé (opcional)
    add_qrcode_footer(doc, qrcode_text)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--id", default="")
    parser.add_argument("--doca", default="")
    parser.add_argument("--data", default=datetime.now().strftime("%d/%m/%Y"))
    parser.add_argument("--conferente", default="")
    parser.add_argument("--turno", default="")
    parser.add_argument("--motorista", default="")
    parser.add_argument("--placa", default="")
    parser.add_argument("--destino", default="")
    parser.add_argument("--notas", default="", help="lista separada por virgulas; se vazio, le do STDIN (uma por linha)")
    parser.add_argument("--logo", default="", help="caminho do arquivo de logo (png/jpg) opcional")
    parser.add_argument("--saida", default="", help="nome do arquivo (opcional)")
    parser.add_argument("--saida_dir", default=".", help="pasta de saida")
    parser.add_argument("--pdf", action="store_true", help="tambem exporta PDF via docx2pdf (requer MS Word)")
    parser.add_argument("--qrcode", default="", help="texto/URL para QR Code no rodape")
    args = parser.parse_args()

    notas = []
    if args.notas.strip():
        notas = [n.strip() for n in args.notas.split(",")]
    else:
        sys.stderr.write("Digite/bipe as notas (uma por linha). Finalize com Ctrl+D (Linux/Mac) ou Ctrl+Z (Windows) e Enter.\n")
        for line in sys.stdin:
            notas.append(line.strip())

    meta = dict(id=args.id, doca=args.doca, data=args.data, conferente=args.conferente,
                turno=args.turno, motorista=args.motorista, placa=args.placa, destino=args.destino)

    # Nome automático
    data_slug = datetime.now().strftime('%Y-%m-%d')
    rota_slug = (args.destino or '').strip().replace(' ', '_')[:40]
    id_slug = (args.id or datetime.now().strftime('%H%M%S'))
    base_name = f"CHECKLIST_{data_slug}_{rota_slug}_{id_slug}".strip('_')

    out_dir = Path(args.saida_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if args.saida:
        saida_path = Path(args.saida)
        if not saida_path.is_absolute():
            saida_path = out_dir / saida_path
    else:
        saida_path = out_dir / f"{base_name}.docx"

    make_docx(str(saida_path), meta, notas, logo_file=args.logo or None, qrcode_text=(args.qrcode or None))
    print(f"OK! Arquivo gerado: {saida_path}")

    if args.pdf:
        if docx2pdf_convert is None:
            print("[AVISO] Instale docx2pdf para PDF: pip install docx2pdf (requer Microsoft Word no Windows/Mac)")
        else:
            try:
                pdf_path = str(saida_path).rsplit('.',1)[0] + '.pdf'
                docx2pdf_convert(str(saida_path), pdf_path)
                print(f"OK! PDF gerado: {pdf_path}")
            except Exception as e:
                print(f"[ERRO] Falha ao exportar PDF: {e}")


if __name__ == "__main__":
    main()
