#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import io
from datetime import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from pathlib import Path

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

try:
    import qrcode
except Exception:
    qrcode = None

st.set_page_config(page_title="Check-list de Notas – Rodopenha", layout="centered")
st.title("Check-list de Movimentacao – Rodopenha")
st.caption("Preencha os campos, bipar/digitar as notas e clique em Gerar DOCX.")

def add_heading_with_rule(doc, text, logo_path=None):
    table_hdr = doc.add_table(rows=1, cols=2)
    cell_logo = table_hdr.cell(0,0)
    cell_title = table_hdr.cell(0,1)
    if logo_path and os.path.exists(logo_path):
        try:
            run_img = cell_logo.paragraphs[0].add_run()
            run_img.add_picture(logo_path, width=Inches(1.2))
        except Exception:
            pass
    p = cell_title.paragraphs[0]
    run = p.add_run(text); run.bold = True; run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # regra
    p_rule = doc.add_paragraph(); p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p_rule._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8'); bottom.set(qn('w:space'), '1')
    p_bdr.append(bottom); p_pr.append(p_bdr)

def add_qrcode_footer(doc, text):
    if not text or qrcode is None:
        return
    img = qrcode.make(text)
    bio = io.BytesIO(); img.save(bio, format='PNG'); bio.seek(0)
    p = doc.add_paragraph(); run = p.add_run()
    try:
        run.add_picture(bio, width=Inches(0.9))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception:
        pass

def build_docx(meta, notas, logo_path=None, qr_text=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    add_heading_with_rule(doc, "CHECK-LIST DE MOVIMENTACAO – RODOPENHA TRANSPORTES", logo_path=logo_path)

    labels = [
        ("ID:", meta.get("id",""), "Data:", meta.get("data", datetime.now().strftime("%d/%m/%Y"))),
        ("Doca:", meta.get("doca",""), "Turno:", meta.get("turno","")),
        ("Conferente:", meta.get("conferente",""), "Motorista:", meta.get("motorista","")),
        ("Placa:", meta.get("placa",""), "Destino/Rota:", meta.get("destino","")),
    ]
    table = doc.add_table(rows=4, cols=4); table.style = "Table Grid"
    for r, row in enumerate(labels):
        for c in range(0, 4, 2):
            cell_label = table.cell(r, c); cell_value = table.cell(r, c+1)
            run_l = cell_label.paragraphs[0].add_run(row[c]); run_l.bold = True
            cell_value.paragraphs[0].add_run(row[c+1])

    doc.add_paragraph("")
    p = doc.add_paragraph(); run = p.add_run("Notas Fiscais / Chaves (bipar ou digitar):"); run.bold = True

    notas = [n.strip() for n in notas if n.strip()]
    total = len(notas)
    if total == 0:
        doc.add_paragraph("—")
    else:
        cols = 2
        rows = (total + cols - 1)//cols
        t2 = doc.add_table(rows=rows, cols=cols); t2.style = "Light Grid"
        for idx, nf in enumerate(notas):
            r = idx % rows; c = idx // rows
            t2.cell(r, c).text = nf

    doc.add_paragraph("")
    t3 = doc.add_table(rows=3, cols=2); t3.style = "Table Grid"
    t3.cell(0,0).text = "Total de Notas Recebidas:"; t3.cell(0,1).text = str(total)
    t3.cell(1,0).text = "Total de Notas Expedidas:"; t3.cell(1,1).text = ""
    t3.cell(2,0).text = "Divergencias:"; t3.cell(2,1).text = ""

    doc.add_paragraph("")
    t4 = doc.add_table(rows=2, cols=2); t4.style = "Table Grid"
    t4.cell(0,0).text = "\n\nAssinatura Conferente:"; t4.cell(0,1).text = "\n\nAssinatura Motorista:"
    t4.cell(1,0).text = "Nome legivel:"; t4.cell(1,1).text = "Nome legivel:"
    doc.add_paragraph("")
    foot = doc.add_paragraph("Documento gerado automaticamente para controle interno de movimentacao.")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER; foot.runs[0].font.size = Pt(8)

    add_qrcode_footer(doc, qr_text)
    return doc

with st.sidebar:
    st.subheader("Opções")
    use_logo = st.checkbox("Incluir logo no cabeçalho", value=True)
    logo_choice = "LogoWebRodoPenha-removebg-preview.png"
    logo_path = os.path.join(os.getcwd(), logo_choice) if use_logo else None
    st.markdown('---')
    saida_dir = st.text_input('Pasta de saída', value='.')
    export_pdf = st.checkbox('Exportar PDF (docx2pdf)', value=False)
    qr_text = st.text_input('QR Code (opcional)', value='')

with st.form("dados"):
    col1, col2 = st.columns(2)
    with col1:
        id_seq = st.text_input("ID")
        doca = st.text_input("Doca")
        conferente = st.text_input("Conferente")
        placa = st.text_input("Placa")
    with col2:
        data_str = st.text_input("Data (dd/mm/aaaa)", value=datetime.now().strftime("%d/%m/%Y"))
        turno = st.text_input("Turno (1º / 2º / 3º)")
        motorista = st.text_input("Motorista")
        destino = st.text_input("Destino/Rota")
    notas_str = st.text_area("Notas (uma por linha – pode bipar o codigo e apertar Enter):", height=220)
    submitted = st.form_submit_button("Gerar DOCX")

if submitted:
    notas = notas_str.splitlines()
    meta = dict(id=id_seq, doca=doca, data=data_str, conferente=conferente,
                turno=turno, motorista=motorista, placa=placa, destino=destino)
    doc = build_docx(meta, notas, logo_path=logo_path, qr_text=qr_text if qr_text else None)

    # disponibilizar para download imediato
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    st.success("Documento gerado!")
    st.download_button("Baixar DOCX", data=bio.getvalue(), file_name="CHECKLIST_MOVIMENTACAO.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # salvar em disco com nome padronizado
    try:
        data_slug = datetime.now().strftime('%Y-%m-%d')
        rota_slug = (destino or '').strip().replace(' ', '_')[:40]
        id_slug = (id_seq or datetime.now().strftime('%H%M%S'))
        base_name = f"CHECKLIST_{data_slug}_{rota_slug}_{id_slug}".strip('_')

        out_dir = Path(saida_dir); out_dir.mkdir(parents=True, exist_ok=True)
        docx_path = out_dir / f"{base_name}.docx"
        doc.save(str(docx_path))
        st.success(f"Arquivo salvo: {docx_path}")

        if export_pdf:
            if docx2pdf_convert is None:
                st.warning('docx2pdf nao disponivel. Instale com: pip install docx2pdf (requer MS Word).')
            else:
                try:
                    pdf_path = out_dir / f"{base_name}.pdf"
                    docx2pdf_convert(str(docx_path), str(pdf_path))
                    st.success(f"PDF salvo: {pdf_path}")
                except Exception as e:
                    st.error(f"Falha ao gerar PDF: {e}")
    except Exception as e:
        st.warning(f"Nao foi possivel salvar automaticamente: {e}")

    st.info("Dica: use leitor de codigo de barras USB – ele digita o codigo e pressiona Enter automaticamente.")


